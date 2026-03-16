import sys
import os
import io
from datetime import datetime
from typing import List, Dict, Any
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))


def export_chat_markdown(messages: List[Dict[str, Any]]) -> str:
    """Export chat as Markdown."""
    md_content = f"# Chat Export\n\n*Exported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}*\n\n---\n\n"
    
    for msg in messages:
        role = msg.get("role", "unknown").capitalize()
        content = msg.get("content", "")
        source = msg.get("source", "")
        
        md_content += f"## {role}\n\n"
        if source and role == "Assistant":
            md_content += f"*Source: {source}*\n\n"
        md_content += f"{content}\n\n---\n\n"
    
    return md_content


def export_chat_html(messages: List[Dict[str, Any]]) -> str:
    """Export chat as HTML."""
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Chat Export</title>
        <style>
            body {{ font-family: Arial, sans-serif; margin: 20px; background: #f5f5f5; }}
            .message {{ background: white; padding: 15px; margin: 10px 0; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1); }}
            .user {{ border-left: 4px solid #60a5fa; }}
            .assistant {{ border-left: 4px solid #4ade80; }}
            .source {{ color: #666; font-size: 12px; margin-top: 5px; }}
            .timestamp {{ color: #999; font-size: 11px; }}
        </style>
    </head>
    <body>
        <h1>Chat Export</h1>
        <p class="timestamp">Exported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        <hr>
    """
    
    for msg in messages:
        role = msg.get("role", "unknown")
        content = msg.get("content", "")
        source = msg.get("source", "")
        
        html_content += f'<div class="message {role}">'
        html_content += f'<strong>{role.capitalize()}</strong>'
        if source and role == "assistant":
            html_content += f'<div class="source">Source: {source}</div>'
        html_content += f'<div>{content}</div>'
        html_content += '</div>'
    
    html_content += """
    </body>
    </html>
    """
    
    return html_content


def export_chat_pdf(messages: List[Dict[str, Any]]) -> bytes:
    """Export chat as PDF."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            textColor='#333333'
        )
        story.append(Paragraph("Chat Export", title_style))
        story.append(Spacer(1, 0.2*inch))
        story.append(Paragraph(f"Exported on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Messages
        for msg in messages:
            role = msg.get("role", "unknown").capitalize()
            content = msg.get("content", "")
            source = msg.get("source", "")
            
            role_style = ParagraphStyle(
                'RoleStyle',
                parent=styles['Heading2'],
                fontSize=12,
                textColor='#60a5fa' if role == "User" else '#4ade80'
            )
            
            story.append(Paragraph(f"<b>{role}</b>", role_style))
            if source and role == "Assistant":
                story.append(Paragraph(f"<i>Source: {source}</i>", styles['Italic']))
            story.append(Paragraph(content.replace('\n', '<br/>'), styles['Normal']))
            story.append(Spacer(1, 0.2*inch))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.read()
    
    except Exception as e:
        # Fallback to empty PDF
        return b""


def export_chat_word(messages: List[Dict[str, Any]]) -> bytes:
    """Export chat as Word document."""
    try:
        from docx import Document
        from docx.shared import Inches
        from io import BytesIO
        
        doc = Document()
        doc.add_heading('Chat Export', 0)
        doc.add_paragraph(f'Exported on {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        doc.add_paragraph('')
        
        for msg in messages:
            role = msg.get("role", "unknown").capitalize()
            content = msg.get("content", "")
            source = msg.get("source", "")
            
            doc.add_heading(role, level=2)
            if source and role == "Assistant":
                doc.add_paragraph(f'Source: {source}', style='Intense Quote')
            doc.add_paragraph(content)
            doc.add_paragraph('')
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    except Exception as e:
        return b""

